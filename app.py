import os
import re
import time
from flask import Flask, request, jsonify, send_file
from google.genai import Client
import json
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

app = Flask(__name__, static_folder='.', static_url_path='')

# --- CONFIGURACIÓN DE GEMINI API (¡VERSIÓN SEGURA!) ---
# [CAMBIO DE SEGURIDAD] Uso exlusivo de variable de entorno (Render/Local)
API_KEY = os.environ.get("GEMINI_API_KEY")

if not API_KEY:
    # [VALIDACIÓN] Error bloqueante si no existe la clave
    raise ValueError("FATAL: La variable de entorno GEMINI_API_KEY no está definida. Configurela en el dashboard de Render o en su entorno local.")


try:
    client = Client(api_key=API_KEY)
except Exception as e:
    print(f"ERROR: No se pudo crear el cliente Gemini. Detalle: {e}")

# Variable global para almacenar la última sesión generada
ultima_sesion = None
import threading
suggest_lock = threading.Lock()


def extract_json_from_text(text):
    """
    Intenta extraer un objeto JSON válido de un texto que puede contener Markdown.
    """
    # 1. Intentar encontrar bloque de código JSON ```json ... ```
    match = re.search(r'```json\s*(\{.*?\})\s*```', text, re.DOTALL)
    if match:
        return match.group(1)
    
    # 2. Intentar parsear todo el texto directamente (por si ya es JSON limpio)
    try:
        json.loads(text)
        return text
    except json.JSONDecodeError:
        pass

    # 3. Intentar encontrar cualquier bloque entre llaves {}
    # Buscamos la primera llave { y la última llave }
    match = re.search(r'(\{.*\})', text, re.DOTALL)
    if match:
        return match.group(1)
        
    # 4. Si no hay marcadores, intentar el texto completo (limpiando espacios)
    return text.strip()


def generate_with_retry(model_name, prompt, retries=5, delay=3):
    """
    Intenta generar contenido con reintentos automáticos en caso de error 503 (Sobrecarga) o 429 (Resource Exhausted).
    """
    for attempt in range(retries):
        try:
            response = client.models.generate_content(
                model=model_name,
                contents=prompt
            )
            return response
        except Exception as e:
            error_str = str(e)
            # Detectar error 503, "overloaded" o 429 "ResourceExhausted"
            is_overloaded = "503" in error_str or "overloaded" in error_str.lower()
            is_resource_exhausted = "429" in error_str or "ResourceExhausted" in error_str or "RESOURCE_EXHAUSTED" in error_str

            if is_overloaded or is_resource_exhausted:
                if attempt < retries - 1:
                    wait_time = delay * (2 ** attempt) # Backoff exponencial: 2s, 4s...
                    print(f"ALERTA: Modelo ocupado (429/503). Reintentando en {wait_time}s... (Intento {attempt + 1}/{retries})")
                    time.sleep(wait_time)
                    continue
            
            # Si es el último intento y sigue fallando con 429/503, lanzamos un error que podamos capturar
            if attempt == retries - 1 and (is_overloaded or is_resource_exhausted):
                raise Exception("SERVICE_UNAVAILABLE_429")
                
            raise e



@app.route('/')
def index():
    try:
        with open('index.html', 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        return "Error: No se encuentra el archivo index.html.", 500


@app.route('/suggest', methods=['POST'])
def generar_sugerencias():
    
    datos = request.json
    campo = datos.get("campo")
    tema = datos.get("tema")
    nivel = datos.get("nivel")
    grado = datos.get("grado")
    area = datos.get("area")

    # Validar datos estrictamente
    if not all([campo, tema, nivel, grado, area]):
        return jsonify({"error": "Faltan datos requeridos (Tema, Nivel, Grado o Área)."}), 400
    # -------------------------------------------------------------
    # FALLBACK PROFESIONAL (Respaldo estático cuando falla la AI)
    # -------------------------------------------------------------
    FALLBACK_SUGGESTIONS = {
        "Inicial": {
            "desempeno": ["Participa en conversaciones espontáneas.", "Explora materiales con sus sentidos.", "Reconoce partes de su cuerpo.", "Expresa sus emociones verbal y no verbalmente."]
        },
        "Primaria": {
            "desempeno": ["Recupera información explícita de textos orales.", "Explica el tema y el propósito comunicativo.", "Deduce características implícitas de personas y personajes.", "Adecúa el texto a la situación comunicativa."]
        },
        "Secundaria": {
            "desempeno": ["Identifica información explícita, relevante y complementaria.", "Infiere e interpreta información del texto escrito.", "Justifica su posición sobre textos leídos.", "Evalúa el uso del lenguaje y la intención del autor."]
        }
    }

    try:
        prompt = f"""
Eres experto en el Currículo MINEDU.
Con los datos:
Tema: {tema}
Nivel: {nivel}
Grado: {grado}
Área: {area}

Genera SOLO un JSON válido:
{{
    "{campo}_sugerencias": [
        "Opción 1",
        "Opción 2",
        "Opción 3",
        "Opción 4"
    ]
}}
"""

        try:
            response = generate_with_retry("gemini-2.5-flash-lite", prompt)
            
            texto = response.text.strip()
            json_text = extract_json_from_text(texto)
            data = json.loads(json_text)
            key = campo + "_sugerencias"
            
            if key in data:
                return jsonify({"sugerencias": data[key]})
            
            # Si la key no está, intentar buscar cualquier lista
            first_key = list(data.keys())[0] if data.keys() else None
            if first_key and isinstance(data[first_key], list):
                return jsonify({"sugerencias": data[first_key]})

        except Exception as e:
            # Si falla la IA (429, 500, Json error...), usamos FALLBACK
            print(f"⚠️ ERROR IA ({e}) -> USANDO FALLBACK LOCAL")
            
            # Determinar fallback según nivel
            nivel_key = nivel if nivel in FALLBACK_SUGGESTIONS else "Primaria"
            campo_key = "desempeno" # Por ahora solo tenemos fallback robusto para desempeno
            
            if campo_key in FALLBACK_SUGGESTIONS[nivel_key]:
                return jsonify({"sugerencias": FALLBACK_SUGGESTIONS[nivel_key][campo_key]})
            
            # Si no hay fallback específico, devolver genérico
            return jsonify({"sugerencias": ["Opción sugerida estándar 1 (IA ocupada)", "Opción sugerida estándar 2 (IA ocupada)"]})

        return jsonify({"error": "No se pudieron generar sugerencias."}), 500

    except Exception as e:
        print(f"ERROR CRITICO EN SUGGEST: {e}")
        return jsonify({"error": f"Error interno: {str(e)}"}), 500



@app.route('/generate', methods=['POST'])
def generar_sesion():
    global ultima_sesion

    try:
        datos = request.json

        # OBTENER TODOS LOS DATOS DEL FORMULARIO
        tema = datos.get('tema', 'Tema no especificado')
        nivel = datos.get('nivel', 'N/A')
        grado = datos.get('grado', 'N/A')
        area = datos.get('area', 'N/A')
        competencia = datos.get('competencia', 'N/A')
        capacidad = datos.get('capacidad', 'N/A')
        desempeno = datos.get('desempeno', 'N/A')
        comp_transversal = datos.get('comp_transversal', 'N/A')
        cap_transversal = datos.get('cap_transversal', 'N/A')
        enfoque = datos.get('enfoque', 'N/A')
        valor = datos.get('valor', 'N/A')
        
        try:
            tiempo_total_user = int(datos.get('tiempo', 90))
        except ValueError:
            tiempo_total_user = 90

        tipo_sesion = datos.get('tipo_sesion', 'Detallada')
        
        # --- PROMPT MAESTRO OPTIMIZADO DEL USUARIO ---
        prompt = f"""
Eres un especialista en diseño curricular del MINEDU – Perú.
Actúas como asistente pedagógico del docente.

REGLAS ABSOLUTAS (NO ROMPER):
1. NO cambies el título de la sesión.
2. NO inventes ni modifiques competencias, capacidades, desempeños, área, grado o nivel.
3. SOLO desarrolla pedagógicamente la sesión.
4. RESPONDE ÚNICAMENTE con un JSON válido (sin texto adicional, sin markdown).
5. NO uses ```json ni explicaciones.

DATOS DEFINIDOS POR EL DOCENTE (NO MODIFICAR):
- Título de la sesión (tema): {tema}
- Nivel: {nivel}
- Grado: {grado}
- Área: {area}
- Competencia: {competencia}
- Capacidad: {capacidad}
- Desempeño: {desempeno}
- Competencia transversal: {comp_transversal}
- Capacidad transversal: {cap_transversal}
- Enfoque transversal: {enfoque}
- Valor: {valor}
- Tiempo total: {tiempo_total_user} minutos
- Tipo de sesión: {tipo_sesion}

INSTRUCCIÓN PEDAGÓGICA:
- Si Tipo de sesión = "Resumida":
  - Redacta actividades breves, claras y directas.
  - Usa listas cortas o párrafos concisos.
- Si Tipo de sesión = "Detallada":
  - Redacta actividades extensas, explicativas y narrativas.
  - Incluye acciones del docente, del estudiante y preguntas orientadoras.

SALIDA OBLIGATORIA (JSON EXACTO):
{{
  "titulo_sesion": "{tema}",
  "proposito": "Describe el propósito de aprendizaje de forma clara y alineada al desempeño.",
  "evidencia": "Describe brevemente qué producto o actuación demostrará el aprendizaje.",
  "estandar_aprendizaje": "Texto completo del Estándar de Aprendizaje del ciclo correspondiente para la competencia seleccionada.",
  "datos_adicionales": {{
    "competencia_transversal": "{comp_transversal}",
    "capacidad_transversal": "{cap_transversal}",
    "enfoque_transversal": "{enfoque}",
    "valor_asociado": "{valor}",
    "tiempo_total": "{tiempo_total_user} minutos"
  }},
  "criterios_evaluacion": [
    "Criterio observable alineado al desempeño.",
    "Criterio medible relacionado con la competencia."
  ],
  "secuencia_didactica": {{
    "inicio": "Describe las actividades de inicio respetando el tipo de sesión.",
    "desarrollo": "Describe las actividades de desarrollo respetando el tipo de sesión.",
    "cierre": "Describe las actividades de cierre respetando el tipo de sesión."
  }}
}}
"""

        # Llamada a la API de Gemini con Retry
        response = generate_with_retry(
            model_name='gemini-2.5-flash-lite',
            prompt=prompt
        )

        response_text = response.text.strip()

        # Extracción robusta de JSON
        json_text = extract_json_from_text(response_text)

        sesion_data = json.loads(json_text)

        # Guardar la sesión para la descarga
        ultima_sesion = {
            'datos_form': datos,
            'sesion': sesion_data
        }

        # Devolver el JSON de la sesión
        return jsonify({'sesion': sesion_data})

    except json.JSONDecodeError as e:
        error_msg = f"Error de formato JSON. La IA respondió: {response_text[:200]}..."
        print(f"TRACEBACK: {str(e)}")
        print(f"RAW RESPONSE TEXT:\n{response_text}") # Log clave para depuración
        return jsonify({'error': error_msg}), 500

    except Exception as e:
        error_msg = str(e)
        if "SERVICE_UNAVAILABLE_429" in error_msg:
             return jsonify({'error': 'La IA está ocupada en este momento. Por favor espera unos segundos.'}), 429

        print(f"TRACEBACK: {error_msg}")
        return jsonify({'error': f"Error al generar la sesión: {error_msg}"}), 500


@app.route('/generate_ept_structure', methods=['POST'])
def generate_ept_structure():
    """
    Genera un CONJUNTO de opciones curriculares para EPT (Competencias, Capacidades, Desempeños)
    para que el docente elija.
    """
    datos = request.json
    nivel = datos.get("nivel", "Secundaria")
    grado = datos.get("grado", "Grado no especificado")
    especialidad = datos.get("especialidad", "No especificada")
    tema = datos.get("tema", "Tema no especificado")

    # Validar que sea EPT (aunque el frontend debe controlar esto) y que haya especialidad
    if not especialidad:
         return jsonify({"error": "La especialidad es obligatoria para EPT."}), 400

    prompt = f"""
Actúa como especialista del Ministerio de Educación del Perú (MINEDU),
experto en Educación para el Trabajo (EPT).

Contexto:
El usuario ha seleccionado:
- Área: Educación para el Trabajo (EPT)
- Especialidad: {especialidad}
- Nivel: {nivel}
- Grado/Año: {grado}
- Tema: {tema}

Objetivo:
Generar un conjunto de OPCIONES curriculares
para que el docente pueda ELEGIR,
no para asignar automáticamente.

Reglas estrictas:
1. NO completes campos automáticamente.
2. Genera OPCIONES, no decisiones finales.
3. Devuelve TODO en una sola respuesta.
4. Contenido alineado al enfoque del CNEB – EPT Perú.
5. Lenguaje técnico, claro y docente.
6. Optimizado para carga rápida (máx. 1 llamada IA).

Formato de salida (JSON puro):

{{
  "competencias": [
    {{
      "nombre": "Nombre de la Competencia 1",
      "capacidades": ["Capacidad 1.1", "Capacidad 1.2", "Capacidad 1.3"],
      "desempenos": ["Desempeño 1.1", "Desempeño 1.2", "Desempeño 1.3", "Desempeño 1.4"]
    }},
    {{
      "nombre": "Nombre de la Competencia 2",
      "capacidades": ["Capacidad 2.1", "Capacidad 2.2"],
      "desempenos": ["Desempeño 2.1", "Desempeño 2.2", "Desempeño 2.3"]
    }}
  ]
}}

Cantidad:
- 3 a 4 competencias
- 3 a 5 capacidades por competencia
- 4 a 6 desempeños por competencia

Genera solo el JSON.
"""
    try:
        # Usamos generate_with_retry existente
        response = generate_with_retry("gemini-2.5-flash-lite", prompt)
        
        texto = response.text.strip()
        json_text = extract_json_from_text(texto)
        data = json.loads(json_text)
        
        # Validación básica de estructura
        if "competencias" not in data or not isinstance(data["competencias"], list):
             # Intento de corrección si la IA devolvió lista directa
             if isinstance(data, list):
                 data = {"competencias": data}
             else:
                 raise ValueError("Estructura JSON inválida: Falta clave 'competencias'")

        return jsonify(data)

    except Exception as e:
        print(f"ERROR EPT: {str(e)}")
        return jsonify({"error": str(e)}), 500


@app.route('/download')
def descargar_word():
    global ultima_sesion

    if not ultima_sesion:
        return jsonify({'error': 'No hay sesión generada'}), 400

    try:
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx import Document
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from io import BytesIO

        # --- RECURSOS COMPARTIDOS (Helpers) ---
        def shade_cell(cell, shade_color):
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), shade_color)
            cell._element.get_or_add_tcPr().append(shading_elm)

        def add_bold_centered_text(cell, text, font_size=10, shade_color=None):
            if shade_color:
                shade_cell(cell, shade_color)
            if cell.paragraphs:
                p = cell.paragraphs[0]
            else:
                p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(text) if text else '-')
            run.bold = True
            run.font.size = Pt(font_size)
            run.font.name = 'Arial'

        def add_bold_text(cell, text, font_size=10, shade_color=None, align='left'):
            if shade_color:
                shade_cell(cell, shade_color)
            if cell.paragraphs:
                p = cell.paragraphs[0]
            else:
                p = cell.add_paragraph()
            if align == 'center':
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'justify':
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif align == 'right':
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(str(text) if text else '-')
            run.bold = True
            run.font.size = Pt(font_size)
            run.font.name = 'Arial'

        def add_normal_text(cell, text, font_size=10, align='left'):
            if cell.paragraphs:
                p = cell.paragraphs[0]
            else:
                p = cell.add_paragraph()
            if align == 'center':
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'justify':
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif align == 'right':
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(str(text) if text else '-')
            run.font.size = Pt(font_size)
            run.font.name = 'Arial'

        def format_criterios(criterios_list):
            if isinstance(criterios_list, list):
                return "\n".join([f"• {c}" for c in criterios_list])
            return str(criterios_list)

        # --- GENERADORES ESPECÍFICOS ---

        def generar_docx_secundaria(doc, datos, sesion):
            # ... Data Preparation ...
            datos_adicionales = sesion.get('datos_adicionales', {})
            try:
                tiempo_total = int(str(datos_adicionales.get('tiempo_total', '90')).split(' ')[0])
            except (ValueError, IndexError):
                tiempo_total = 90
            
            inicio_min = round(tiempo_total * 0.2)
            desarrollo_min = round(tiempo_total * 0.6)
            cierre_min = round(tiempo_total * 0.2)
            
            # Variables de tiempo para el doc
            ini_str = f"{inicio_min}'"
            des_str = f"{desarrollo_min}'"
            cie_str = f"{cierre_min}'"

            # ... Datos Administrativos ...
            dre = datos.get('dre', 'San Martín')
            ugel = datos.get('ugel', 'San Martín')
            ie = datos.get('ie', 'N/A')
            distrito = datos.get('distrito', 'Tarapoto')
            seccion = datos.get('seccion', 'A, B, C, D')
            ciclo = datos.get('ciclo', 'VI') # Secundaria suele ser VI o VII
            director = datos.get('director', 'N/A')
            docente = datos.get('docente', 'N/A')
            fecha = datos.get('fecha', 'N/A')
            duracion = datos.get('duracion', '90\'')


            # 1. TÍTULO PRINCIPAL
            titulo_principal = doc.add_paragraph()
            titulo_principal.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_titulo = titulo_principal.add_run('SESIÓN DE APRENDIZAJE\nINNOVACIÓN PEDAGÓGICA')
            run_titulo.bold = True
            run_titulo.font.size = Pt(14)
            run_titulo.font.name = 'Arial'
            doc.add_paragraph()

            # 2. TÍTULO SESIÓN
            tabla_titulo = doc.add_table(rows=1, cols=1)
            tabla_titulo.style = 'Table Grid'
            add_bold_centered_text(tabla_titulo.cell(0, 0), datos.get('tema', ''), 11)
            doc.add_paragraph()

            # 3. DATOS INFORMATIVOS
            p_datos = doc.add_paragraph()
            p_datos.add_run('☰ ').bold = True
            p_datos.add_run('DATOS INFORMATIVOS:').bold = True
            
            tabla_info = doc.add_table(rows=5, cols=8)
            tabla_info.style = 'Table Grid'
            
            add_bold_text(tabla_info.cell(0, 0), 'DRE', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(0, 1), dre, 9)
            add_bold_text(tabla_info.cell(0, 2), 'UGEL', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(0, 3).merge(tabla_info.cell(0, 7)), ugel, 9)
            
            add_bold_text(tabla_info.cell(1, 0), 'Institución Educativa', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(1, 1).merge(tabla_info.cell(1, 3)), ie, 9)
            add_bold_text(tabla_info.cell(1, 4), 'Distrito', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(1, 5).merge(tabla_info.cell(1, 7)), distrito, 9)
            
            add_bold_text(tabla_info.cell(2, 0), 'Área curricular', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(2, 1), datos.get('area', 'N/A'), 9)
            add_bold_text(tabla_info.cell(2, 2), 'Grado', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(2, 3), datos.get('grado', 'N/A'), 9)
            add_bold_text(tabla_info.cell(2, 4), 'Sección', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(2, 5), seccion, 9)
            add_bold_text(tabla_info.cell(2, 6), 'Duración', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(2, 7), duracion, 9)
            
            add_bold_text(tabla_info.cell(3, 0), 'Ciclo', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(3, 1), ciclo, 9)
            add_bold_text(tabla_info.cell(3, 2), 'Fecha', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(3, 3).merge(tabla_info.cell(3, 5)), fecha, 9)
            add_bold_text(tabla_info.cell(3, 6), 'Director(a)', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(3, 7), director, 9)
            
            add_bold_text(tabla_info.cell(4, 0), 'Docente', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(4, 1).merge(tabla_info.cell(4, 7)), docente, 9)
            doc.add_paragraph()

            # 4. PROPÓSITO
            p_proposito = doc.add_paragraph()
            p_proposito.add_run('☰ ').bold = True
            p_proposito.add_run('PROPÓSITO DE APRENDIZAJE').bold = True
            
            tabla_proposito = doc.add_table(rows=2, cols=5)
            tabla_proposito.style = 'Table Grid'
            add_bold_centered_text(tabla_proposito.cell(0, 0), 'COMPETENCIA', 9, 'D9D9D9')
            add_bold_centered_text(tabla_proposito.cell(0, 1), 'CAPACIDAD', 9, 'D9D9D9')
            add_bold_centered_text(tabla_proposito.cell(0, 2), 'DESEMPEÑOS\nPRECISADOS', 9, 'D9D9D9')
            add_bold_centered_text(tabla_proposito.cell(0, 3), 'CRITERIOS DE EVALUACIÓN', 9, 'D9D9D9')
            add_bold_centered_text(tabla_proposito.cell(0, 4), 'INSTRUMENTO', 9, 'D9D9D9')
            
            row = tabla_proposito.rows[1].cells
            add_normal_text(row[0], datos.get('competencia', ''), 9, 'justify')
            add_normal_text(row[1], datos.get('capacidad', ''), 9, 'justify')
            add_normal_text(row[2], datos.get('desempeno', ''), 9, 'justify')
            add_normal_text(row[3], format_criterios(sesion.get('criterios_evaluacion', [])), 9, 'justify')
            add_normal_text(row[4], 'Lista de cotejo', 9, 'center')
            
            comp_trans = datos.get('comp_transversal')
            if comp_trans and comp_trans not in ['N/A', '']:
                row_t = tabla_proposito.add_row().cells
                add_normal_text(row_t[0], comp_trans, 9, 'justify')
                add_normal_text(row_t[1], datos.get('cap_transversal', ''), 9, 'justify')
                add_normal_text(row_t[2], 'Se desenvuelve...', 9, 'justify')
                add_normal_text(row_t[3], 'Observación', 9, 'justify')
                add_normal_text(row_t[4], 'Ficha', 9, 'center')
            doc.add_paragraph()

            # 5. ENFOQUES
            p_enfoque = doc.add_paragraph()
            p_enfoque.add_run('☰ ').bold = True
            p_enfoque.add_run('ENFOQUES TRANSVERSALES').bold = True
            tabla_enfoques = doc.add_table(rows=2, cols=4)
            tabla_enfoques.style = 'Table Grid'
            add_bold_centered_text(tabla_enfoques.cell(0, 0), 'ENFOQUE', 9, 'D9D9D9')
            add_bold_centered_text(tabla_enfoques.cell(0, 1), 'VALORES', 9, 'D9D9D9')
            add_bold_centered_text(tabla_enfoques.cell(0, 2), 'ACTITUDES', 9, 'D9D9D9')
            add_bold_centered_text(tabla_enfoques.cell(0, 3), 'ACCIONES', 9, 'D9D9D9')
            
            add_normal_text(tabla_enfoques.cell(1, 0), datos.get('enfoque', ''), 9, 'justify')
            add_normal_text(tabla_enfoques.cell(1, 1), datos.get('valor', ''), 9, 'justify')
            add_normal_text(tabla_enfoques.cell(1, 2), 'Actitud de ejemplo.', 9, 'justify')
            add_normal_text(tabla_enfoques.cell(1, 3), 'Acciones observables.', 9, 'justify')
            doc.add_paragraph()

            # 6. SECUENCIA
            p_secuencia = doc.add_paragraph()
            p_secuencia.add_run('☰ ').bold = True
            p_secuencia.add_run('SECUENCIA DIDÁCTICA').bold = True
            tabla_sec = doc.add_table(rows=4, cols=4)
            tabla_sec.style = 'Table Grid'
            add_bold_centered_text(tabla_sec.cell(0, 0), 'MOMENTOS', 10, 'D9D9D9')
            add_bold_centered_text(tabla_sec.cell(0, 1), 'ACTIVIDADES', 10, 'D9D9D9')
            add_bold_centered_text(tabla_sec.cell(0, 2), 'MATERIALES', 10, 'D9D9D9')
            add_bold_centered_text(tabla_sec.cell(0, 3), 'TIEMPO', 10, 'D9D9D9')
            
            sd = sesion.get('secuencia_didactica', {})
            # Inicio
            add_bold_text(tabla_sec.cell(1, 0), 'MOTIVACIÓN', 9)
            add_normal_text(tabla_sec.cell(1, 1), sd.get('inicio', ''), 9, 'justify')
            add_normal_text(tabla_sec.cell(1, 2), 'Recursos clase', 9)
            add_bold_centered_text(tabla_sec.cell(1, 3), ini_str, 10)
            # Desarrollo
            add_bold_text(tabla_sec.cell(2, 0), 'DESARROLLO', 9)
            add_normal_text(tabla_sec.cell(2, 1), sd.get('desarrollo', ''), 9, 'justify')
            add_normal_text(tabla_sec.cell(2, 2), 'Recursos clase', 9)
            add_bold_centered_text(tabla_sec.cell(2, 3), des_str, 10)
            # Cierre
            add_bold_text(tabla_sec.cell(3, 0), 'CIERRE', 9)
            add_normal_text(tabla_sec.cell(3, 1), sd.get('cierre', ''), 9, 'justify')
            add_normal_text(tabla_sec.cell(3, 2), 'Recursos clase', 9)
            add_bold_centered_text(tabla_sec.cell(3, 3), cie_str, 10)
            
            doc.add_paragraph()
            doc.add_paragraph('• Bibliografía referencial.')
            
            # Firmas
            t_firmas = doc.add_table(rows=1, cols=2)
            t_firmas.rows[0].cells[0].paragraphs[0].add_run('___________________\nV.B. Director').alignment = WD_ALIGN_PARAGRAPH.CENTER
            t_firmas.rows[0].cells[1].paragraphs[0].add_run('___________________\nDocente').alignment = WD_ALIGN_PARAGRAPH.CENTER

        def generar_docx_primaria(doc, datos, sesion):
            # Colores
            YELLOW_HEADER = 'FEF2CC' # Amarillo claro 
            GREEN_HEADER = 'E2EFDA' # Verde claro
            GREEN_BRIGHT = '548235' # Verde fuerte para texto

            # ... Data Preparation ...
            datos_adicionales = sesion.get('datos_adicionales', {})
            try:
                tiempo_total = int(str(datos_adicionales.get('tiempo_total', '90')).split(' ')[0])
            except:
                tiempo_total = 90
            
            # Datos Administrativos
            dre = datos.get('dre', 'San Martín')
            ugel = datos.get('ugel', 'San Martín')
            ie = datos.get('ie', '')
            distrito = datos.get('distrito', '')
            seccion = datos.get('seccion', '')
            ciclo = datos.get('ciclo', 'III/IV/V')
            director = datos.get('director', '')
            docente = datos.get('docente', '')
            fecha = datos.get('fecha', '')
            duracion = datos.get('duracion', '90 min')

            # 1. Título General
            doc.add_paragraph() 
            titulo = doc.add_paragraph()
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = titulo.add_run('SESIÓN DE APRENDIZAJE')
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0, 0, 0)
            doc.add_paragraph()

            # 2. Título Sesión (Recuadro)
            tbl_tit = doc.add_table(rows=1, cols=1)
            tbl_tit.style = 'Table Grid'
            add_bold_centered_text(tbl_tit.cell(0, 0), datos.get('tema', 'TEMA DE SESIÓN'), 12)
            doc.add_paragraph()

            p_inf = doc.add_paragraph()
            p_inf.add_run('☰ ').bold = True
            p_inf.add_run('DATOS INFORMATIVOS:').bold = True
            # run_inf = p_inf.add_run('I. DATOS INFORMATIVOS:') # Old style

            tabla_info = doc.add_table(rows=5, cols=8)
            tabla_info.style = 'Table Grid'
            
            add_bold_text(tabla_info.cell(0, 0), 'DRE', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(0, 1), dre, 9)
            add_bold_text(tabla_info.cell(0, 2), 'UGEL', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(0, 3).merge(tabla_info.cell(0, 7)), ugel, 9)
            
            add_bold_text(tabla_info.cell(1, 0), 'Institución Educativa', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(1, 1).merge(tabla_info.cell(1, 3)), ie, 9)
            add_bold_text(tabla_info.cell(1, 4), 'Distrito', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(1, 5).merge(tabla_info.cell(1, 7)), distrito, 9)
            
            add_bold_text(tabla_info.cell(2, 0), 'Área curricular', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(2, 1), datos.get('area', 'N/A'), 9)
            add_bold_text(tabla_info.cell(2, 2), 'Grado', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(2, 3), datos.get('grado', 'N/A'), 9)
            add_bold_text(tabla_info.cell(2, 4), 'Sección', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(2, 5), seccion, 9)
            add_bold_text(tabla_info.cell(2, 6), 'Duración', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(2, 7), duracion, 9)
            
            add_bold_text(tabla_info.cell(3, 0), 'Ciclo', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(3, 1), ciclo, 9)
            add_bold_text(tabla_info.cell(3, 2), 'Fecha', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(3, 3).merge(tabla_info.cell(3, 5)), fecha, 9)
            add_bold_text(tabla_info.cell(3, 6), 'Director(a)', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(3, 7), director, 9)
            
            add_bold_text(tabla_info.cell(4, 0), 'Docente', 9, 'D9D9D9')
            add_normal_text(tabla_info.cell(4, 1).merge(tabla_info.cell(4, 7)), docente, 9)
            
            doc.add_paragraph()

            # 4. PROPÓSITOS DE APRENDIZAJE (Estilo Primaria - Amarillo)
            p_prop = doc.add_paragraph()
            run_prop = p_prop.add_run('II. PROPÓSITOS DE APRENDIZAJE Y EVIDENCIAS DE APRENDIZAJE:')
            run_prop.bold = True
            run_prop.font.color.rgb = RGBColor(255, 0, 0) # Rojo

            t_prop = doc.add_table(rows=2, cols=4)
            t_prop.style = 'Table Grid'
            # Encabezados Amarillos
            add_bold_centered_text(t_prop.cell(0, 0), 'Competencias', 10, YELLOW_HEADER)
            add_bold_centered_text(t_prop.cell(0, 1), 'Capacidades', 10, YELLOW_HEADER)
            add_bold_centered_text(t_prop.cell(0, 2), 'Desempeños', 10, YELLOW_HEADER)
            add_bold_centered_text(t_prop.cell(0, 3), 'Criterios de\nevaluación', 10, YELLOW_HEADER)

            # Contenido
            row_p = t_prop.rows[1].cells
            add_normal_text(row_p[0], datos.get('competencia', ''), 9)
            add_normal_text(row_p[1], datos.get('capacidad', ''), 9)
            add_normal_text(row_p[2], datos.get('desempeno', ''), 9)
            add_normal_text(row_p[3], format_criterios(sesion.get('criterios_evaluacion', [])), 9)

            # TABLA ESTÁNDAR (Amarilla)
            doc.add_paragraph()
            t_stand = doc.add_table(rows=2, cols=2) # 2 filas: Header y Contenido
            t_stand.style = 'Table Grid'
            
            # Header Merged
            cell_header = t_stand.cell(0, 0)
            cell_header.merge(t_stand.cell(0, 1))
            add_bold_text(cell_header, 'ESTÁNDAR DE APRENDIZAJE POR COMPETENCIAS Y GRADOS', 10, YELLOW_HEADER)
            
            # Contenido (Estándar real)
            cell_cont = t_stand.cell(1, 0)
            cell_cont.merge(t_stand.cell(1, 1))
            
            estandar_texto = sesion.get('estandar_aprendizaje', 'No especificado en la sesión generada.')
            add_normal_text(cell_cont, estandar_texto, 9, 'justify')
            
            # Tabla Enfoques (Verde)
            doc.add_paragraph()
            t_enf = doc.add_table(rows=5, cols=3)
            t_enf.style = 'Table Grid'
            # Header color background check image: shows bright Green headers.
            # Lets try to match the image: Green header for "ENFOQUES TRANSVERSALES", "VALORES", "EJEMPLOS"
            add_bold_text(t_enf.cell(0, 0), 'ENFOQUES TRANSVERSALES', 9, '548235')
            add_bold_text(t_enf.cell(0, 1), 'VALORES', 9, '548235')
            add_bold_text(t_enf.cell(0, 2), 'EJEMPLOS', 9, '548235')
            # Shading header cells green
            shade_cell(t_enf.cell(0, 0), '70AD47')
            shade_cell(t_enf.cell(0, 1), '70AD47')
            shade_cell(t_enf.cell(0, 2), '70AD47')
            
            # Contenido Enfoques
            add_normal_text(t_enf.cell(1, 0), datos.get('enfoque', ''), 9)
            add_normal_text(t_enf.cell(1, 1), datos.get('valor', ''), 9)
            add_normal_text(t_enf.cell(1, 2), 'Ejemplo observable...', 9)

            # Competencias Transversales (Yellow Headers separate)
            doc.add_paragraph()
            
            # Comp Transversal 1
            t_ct1 = doc.add_table(rows=2, cols=2)
            t_ct1.style = 'Table Grid'
            add_bold_centered_text(t_ct1.cell(0, 0), 'Competencia transversal', 9, YELLOW_HEADER)
            add_bold_centered_text(t_ct1.cell(0, 1), 'Capacidades Transversales', 9, YELLOW_HEADER)
            comp_t = datos.get('comp_transversal', '')
            cap_t = datos.get('cap_transversal', '')
            add_normal_text(t_ct1.cell(1, 0), comp_t if comp_t else 'Se desenvuelve en entornos virtuales...', 9)
            add_normal_text(t_ct1.cell(1, 1), cap_t if cap_t else 'Personaliza entornos...', 9)

            doc.add_paragraph()
            
            # Comp Transversal 2 (Gestiona su aprendizaje - Default en primaria muchas veces)
            t_ct2 = doc.add_table(rows=2, cols=2)
            t_ct2.style = 'Table Grid'
            add_bold_centered_text(t_ct2.cell(0, 0), 'Competencia transversal', 9, YELLOW_HEADER)
            add_bold_centered_text(t_ct2.cell(0, 1), 'Capacidades Transversales', 9, YELLOW_HEADER)
            add_normal_text(t_ct2.cell(1, 0), 'Gestiona su aprendizaje de manera autónoma', 9)
            add_normal_text(t_ct2.cell(1, 1), 'Define metas de aprendizaje', 9)


            # III. PREPARACIÓN DE LA SESIÓN
            doc.add_paragraph()
            p_prep = doc.add_paragraph()
            run_prep = p_prep.add_run('III. PREPARACIÓN DE LA SESIÓN')
            run_prep.bold = True
            run_prep.font.color.rgb = RGBColor(255, 0, 0)

            t_prep = doc.add_table(rows=2, cols=2)
            t_prep.style = 'Table Grid'
            add_bold_centered_text(t_prep.cell(0, 0), '¿Qué se debe hacer antes de la sesión?', 10, YELLOW_HEADER)
            add_bold_centered_text(t_prep.cell(0, 1), '¿Qué recursos o materiales utilizarán en la sesión?', 10, YELLOW_HEADER)
            add_normal_text(t_prep.cell(1, 0), 'Preparar fichas, revisar materiales.', 10)
            add_normal_text(t_prep.cell(1, 1), 'Plumones, papelógrafos, fichas.', 10)

            # IV. MOMENTOS
            doc.add_paragraph()
            p_mom = doc.add_paragraph()
            run_mom = p_mom.add_run('IV. MOMENTOS DE LA SESIÓN')
            run_mom.bold = True
            run_mom.font.color.rgb = RGBColor(255, 0, 0)
            
            sd = sesion.get('secuencia_didactica', {})

            # Función helper para momento CON CUADRO NEGRO (Table Grid)
            def add_moment_table(moment_name, time_val, content):
                # Tabla 2 filas: 
                # Fila 1: Título (Verde) | Tiempo (Verde)
                # Fila 2: Contenido (Blanco, con bordes negros)
                
                t = doc.add_table(rows=2, cols=2)
                t.style = 'Table Grid' 
                
                # HEADER (Verde)
                # Celda 1: Nombre
                c_name = t.cell(0, 0)
                shade_cell(c_name, 'E2EFDA') # Verde claro header
                p1 = c_name.paragraphs[0]
                run1 = p1.add_run(moment_name)
                run1.bold = True
                run1.font.size = Pt(10)

                # Celda 2: Tiempo
                c_time = t.cell(0, 1)
                shade_cell(c_time, 'E2EFDA') # Verde claro header
                p2 = c_time.paragraphs[0]
                p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run2 = p2.add_run(f"Tiempo aproximado: {time_val} min")
                run2.font.size = Pt(9)
                
                # CONTENIDO (Merged)
                c_content = t.cell(1, 0)
                c_content.merge(t.cell(1, 1))
                add_normal_text(c_content, content, 10, 'justify')

                doc.add_paragraph() # Espacio entre tablas

            add_moment_table('INICIO', '15', sd.get('inicio', ''))
            add_moment_table('DESARROLLO', '60', sd.get('desarrollo', ''))
            add_moment_table('CIERRE', '15', sd.get('cierre', ''))

            # V. RECURSOS / BIBLIOGRAFÍA
            p_rec = doc.add_paragraph()
            run_rec = p_rec.add_run('V. RECURSOS Y BIBLIOGRAFÍA')
            run_rec.bold = True
            run_rec.font.color.rgb = RGBColor(255, 0, 0) # Rojo
            
            t_rec = doc.add_table(rows=1, cols=1)
            t_rec.style = 'Table Grid'
            c_rec = t_rec.cell(0,0)
            
            recursos = sesion.get('recursos_virtuales', [])
            if isinstance(recursos, list):
                recursos_txt = "\n".join([f"• {r}" for r in recursos])
            else:
                recursos_txt = str(recursos)
                
            add_normal_text(c_rec, recursos_txt if recursos_txt else "Bibliografía del MED.", 9)

            # Footer / Docente Ref

            # Footer / Docente Ref
            doc.add_paragraph()
            t_foo = doc.add_table(rows=1, cols=1)
            t_foo.align = WD_ALIGN_PARAGRAPH.CENTER
            p_foo = t_foo.cell(0,0).paragraphs[0]
            p_foo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_foo.add_run('__________________________\nDOCENTE')


        # --- MAIN LOGIC ---
        doc = Document()
        # Margenes
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)

        sesion = ultima_sesion['sesion']
        datos = ultima_sesion['datos_form']
        
        nivel = datos.get('nivel', 'Secundaria')
        
        # Lógica de selección
        if nivel == 'Primaria':
            generar_docx_primaria(doc, datos, sesion)
        else:
            generar_docx_secundaria(doc, datos, sesion)

        # Guardar
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        filename = f"Sesion_{nivel}_{datos.get('grado','').replace(' ','_')}.docx"
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Error al crear el documento Word: {str(e)}'}), 500
if __name__ == '__main__':
    PORT = int(os.environ.get('PORT', 5010))
    HOST = '127.0.0.1'
    app.run(debug=True, host=HOST, port=PORT)

